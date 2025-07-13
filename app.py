from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import tempfile
import shutil
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
import re
import unicodedata

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
ALLOWED_EXTENSIONS = {'docx'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return text
    
    # Remove null characters and normalize
    text = text.replace('\x00', '')
    text = unicodedata.normalize('NFKC', text)
    
    # Remove invalid characters
    text = ''.join(char for char in text if ord(char) < 65536)
    
    return text.strip()

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file with proper handling"""
    try:
        doc = Document(docx_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)
        
        return '\n\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""

def create_pdf_with_hindi_support(text, output_path):
    """Create PDF with Hindi font support"""
    try:
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Define styles
        styles = getSampleStyleSheet()

        pdfmetrics.registerFont(TTFont('KrutiDev011', 'fonts/KrutiDev011Regular.ttf'))
        
        # Create custom style for Hindi text
        hindi_style = ParagraphStyle(
            'HindiStyle',
            parent=styles['Normal'],
            fontName='KrutiDev011',
            fontSize=12,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=6
        )
        
        # Process text and create paragraphs
        story = []
        
        # Split text into paragraphs
        paragraphs = text.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Clean the text
                clean_para = clean_text(para_text)
                
                # Create paragraph
                paragraph = Paragraph(clean_para, hindi_style)
                story.append(paragraph)
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        return True
        
    except Exception as e:
        print(f"Error creating PDF: {e}")
        return False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_docx_to_pdf():
    try:
        # Check if file was uploaded
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        
        # Check if file was selected
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file type
        if not allowed_file(file.filename):
            return jsonify({'error': 'Only .docx files are allowed'}), 400
        
        # Secure filename
        filename = secure_filename(file.filename)
        
        # Create temporary file path
        temp_input_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{filename}")
        output_filename = filename.rsplit('.', 1)[0] + '.pdf'
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
        
        try:
            # Save uploaded file
            file.save(temp_input_path)
            
            # Extract text from DOCX
            text = extract_text_from_docx(temp_input_path)
            
            if not text.strip():
                return jsonify({'error': 'No readable text found in the document'}), 400
            
            print(f"Extracted text sample: {text[:200]}")
            
            # Create PDF
            success = create_pdf_with_hindi_support(text, output_path)
            
            if not success:
                return jsonify({'error': 'Failed to create PDF'}), 500
            
            # Clean up input file
            if os.path.exists(temp_input_path):
                os.remove(temp_input_path)
            
            return jsonify({
                'success': True,
                'filename': output_filename,
                'message': 'File converted successfully',
                'text_sample': text[:100]
            })
            
        except Exception as e:
            # Clean up on error
            if os.path.exists(temp_input_path):
                os.remove(temp_input_path)
            raise e
            
    except Exception as e:
        print(f"Conversion error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_pdf(filename):
    try:
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        # Send file
        response = send_file(
            file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/pdf'
        )
        
        # Clean up file after sending
        def cleanup():
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    print(f"Cleaned up: {filename}")
            except Exception as e:
                print(f"Cleanup error: {e}")
        
        # Schedule cleanup after response
        response.call_on_close(cleanup)
        
        return response
        
    except Exception as e:
        print(f"Download error: {e}")
        return jsonify({'error': 'Error downloading file'}), 500

@app.route('/health')
def health_check():
    return jsonify({'status': 'healthy'})

def cleanup_old_files():
    """Clean up old files"""
    import time
    
    current_time = time.time()
    max_age = 30 * 60  # 30 minutes
    
    # Clean uploads folder
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.isfile(file_path):
            if current_time - os.path.getmtime(file_path) > max_age:
                try:
                    os.remove(file_path)
                    print(f"Cleaned up old file: {filename}")
                except Exception as e:
                    print(f"Error cleaning up {filename}: {e}")
    
    # Clean output folder
    for filename in os.listdir(app.config['OUTPUT_FOLDER']):
        file_path = os.path.join(app.config['OUTPUT_FOLDER'], filename)
        if os.path.isfile(file_path):
            if current_time - os.path.getmtime(file_path) > max_age:
                try:
                    os.remove(file_path)
                    print(f"Cleaned up old output: {filename}")
                except Exception as e:
                    print(f"Error cleaning up {filename}: {e}")

# Schedule cleanup
import threading
import time

def cleanup_scheduler():
    while True:
        cleanup_old_files()
        time.sleep(600)  # Run every 10 minutes

# Start cleanup thread
cleanup_thread = threading.Thread(target=cleanup_scheduler, daemon=True)
cleanup_thread.start()

if __name__ == '__main__':
    print("DOCX to PDF Converter with Hindi support starting...")
    print(f"Upload directory: {app.config['UPLOAD_FOLDER']}")
    print(f"Output directory: {app.config['OUTPUT_FOLDER']}")
    app.run(debug=True, host='0.0.0.0', port=5000) 
